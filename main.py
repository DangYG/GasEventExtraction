#coding=utf-8

import numpy as np
import docx



# doc = docx.Document('D:\\知识图谱代码实现\\BiLSTM-CRF\\燃气事故.docx')
# for p in doc.paragraphs:
#     if p.style.name == 'Normal':
#         print(p.text)#输出读取的文本
#         demo_sent = p.text
#     if demo_sent == '' or demo_sent.isspace():
#         print('See you next time!')
#         break
#     else:
#         demo_sent = list(demo_sent.strip())
#         demo_data = [(demo_sent, ['O'] * len(demo_sent))]
doc = docx.Document('D:\\知识图谱代码实现\\BiLSTM-CRF\\燃气事故.docx')
for p in doc.paragraphs:
    print(type(p.text))#输出读取的文本